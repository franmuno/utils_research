import os
import argparse
import re
import unicodedata
import json
import pandas as pd
import docx  
from docxtpl import DocxTemplate, RichText
from markupsafe import escape

from docx import Document  # Using python-docx to read the docx file

# Sample config to show if not present
SAMPLE_CONFIG = {
    "excel_folder": "xlsx/",
    "template_folder": "template/",
    "tmp_result_folder": "tmp_processing/",
    "master_input_filename": "Master_list_IARC2024_30May_BOA.xlsx",
    "master_output_filename": "master_output_v3.xlsx",
    "template_input_filename": "tpl_resumenes_IARC2024_v6.docx",
    "template_separator_filename": "tpl_separadores_IARC2024_v3.docx",
    "sheet_name": 0,
    "abstract_id_column": "AbstractId",
    "abstract_sessionname_column": "Theme",
    "id_column_name": "AbstractId",

    "session_names_code_dict": {
        "ARs as a component of compound events": "AR1",
        "ARs in past, present, and future climates": "AR2",
        "Environmental and socioeconomic impacts ARs": "AR3",
        "Forecasting of ARs": "AR4",
        "Observing, identification, and monitoring of ARs": "AR5",
        "Physical, dynamic, & microphysic aspects of ARs": "AR6",
        "Role of ARs in the changing Cryosphere": "AR7",
        "Rejected": "R"
    },
}

def show_sample_config():
    print("\nSample Configuration (save this as a config.json file):")
    print(json.dumps(SAMPLE_CONFIG, indent=4))

# Function to extract and display the template fields
def extract_template_fields(template_path):
    # Check if the file exists
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found at '{template_path}'")

    # Try opening the document
    try:
        doc = Document(template_path)
    except Exception as e:
        raise Exception(f"Error opening the .docx file: {e}")
    
    # Extract all text from the document (paragraphs and tables)
    doc_text = ""
    for paragraph in doc.paragraphs:
        doc_text += paragraph.text + "\n"
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text += cell.text + "\n"

    # Regular expression to find all placeholders like {{ variable_name }}
    variable_pattern = re.compile(r"{{\s*(\w+)\s*}}")

    # Find all unique variable names in the template
    fields = set(variable_pattern.findall(doc_text))

    return fields

def main():
    parser = argparse.ArgumentParser(description='Generate documents from a template using configuration.')
    parser.add_argument('-c', '--config', type=str, help='Path to the configuration JSON file', default='config.json')
    args = parser.parse_args()

    config_path = args.config

    # Check if the config file exists
    if not os.path.exists(config_path):
        print(f"\nError: Configuration file '{config_path}' not found.")
        show_sample_config()
        return

    # Load the config file
    with open(config_path, 'r') as file:
        config = json.load(file)

    # Extract config values
    excel_folder = config['excel_folder']
    template_folder = config['template_folder']
    tmp_result_folder = config['tmp_result_folder']
    master_input_filename = config['master_input_filename']
    master_output_filename = config['master_output_filename']
    template_input_filename = config['template_input_filename']
    template_separator_filename = config['template_separator_filename']

    sheet_name = config['sheet_name']
    AbstractId = config['abstractid_column']
    Abstract = config['abstract_column']
    AreaName = config['areaname_column']

    # READ MASTER EXCEL WITH ALL DATA
    df = pd.read_excel(os.path.join(excel_folder, master_input_filename), sheet_name=sheet_name, engine='openpyxl', dtype=str)
    df.dropna(subset=[AbstractId], inplace=True)
    df = df.fillna('')

    # ABSTRACT TEMPLATE FILENAMES
    abstract_template_path_name=os.path.join(template_folder, template_input_filename)
    separator_template_path_name=os.path.join(template_folder, template_separator_filename)

    areacode_dict={}
    try:
        AreaCode = config['areacode_column']
        ##TODOarea_codes_dict = process_area_codes(config, df)  
    except KeyError:
        print("Note: 'areacode_column' not found in the config file. Loading areacode_dict.")
        areacode_dict = config['areacode_dict']
        AreaCode = '__AreaCode'
        df[AreaCode] = df[AreaName].map(areacode_dict).fillna('Unknown')       
        # Generate a dict of distinct AreaCodes
        

    # Rename Area with Included Code
    df[AreaName]= df[AreaCode]+ ' - ' + df[AreaName]


    # Add derived columns
    df['__AbstractIdFill'] = df[AbstractId].apply(lambda x: str(x).zfill(3))
    df['__WordCountSumm'] = df[Abstract].apply(lambda n: len(str(n).split()) if not isinstance(n, float) else 0)


    df.sort_values([AreaCode, '__AbstractIdFill'], ascending=[True, True], inplace=True)

    # Ensure the tmp result folder exists
    os.makedirs(tmp_result_folder, exist_ok=True)
    df.to_excel(os.path.join(excel_folder, master_output_filename), sheet_name='Summary')


    # Extract fields from the main template
    extracted_fields = extract_template_fields(abstract_template_path_name)
    print(f"Fields found in template '{template_input_filename}': {extracted_fields}")

    # Extract fields from the separator template
    ##separator_template_path = os.path.join(template_folder, template_separator_filename)
    ##separator_template_fields = extract_template_fields(separator_template_path)
    ##print(f"Fields found in separator template '{template_separator_filename}': {separator_template_fields}")

    # Iterate through the rows of abstracts
    for _, row in df.iterrows():
        doc = DocxTemplate(abstract_template_path_name)

        context = {field: escape(row[field]) for field in extracted_fields}

        doc.render(context)
        doc.add_page_break()

        final_filename = os.path.join(tmp_result_folder, f"{row[AreaCode]}_{row['__AbstractIdFill']}.docx")
        doc.save(final_filename)
        print(f"Single Abstract Document saved: {final_filename}")

    # Iterate through the areas/themes
    ##extracted_fields_sep = extract_template_fields(separator_template_path_name)
    for areaname, areacode  in areacode_dict.items():
        context = { 
            AreaName: areacode+' - '+ areaname
        }
        doc = DocxTemplate(os.path.join(template_folder, template_separator_filename))
        doc.render(context)
        doc.add_page_break()

        final_filename = os.path.join(tmp_result_folder, f"{areacode}_000.docx")
        doc.save(final_filename)
        print(f"Document saved: {final_filename}")


if __name__ == "__main__":
    main()
